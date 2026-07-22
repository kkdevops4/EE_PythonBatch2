import os
import shutil
from datetime import datetime
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

class WordReport:
    def __init__(self, data, driver_report, ranking):
        self.data = data
        self.driver_report = driver_report.copy()
        self.ranking = ranking
        self.prepare_status_data()
        self.document = Document()
        os.makedirs("reports", exist_ok=True)
        os.makedirs("backup", exist_ok=True)
        self.chart_folder = "reports/temp_charts"
        os.makedirs(self.chart_folder, exist_ok=True)

        style = self.document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        section = self.document.sections[0]
        section.header.paragraphs[0].text = "Driver Monitoring System Analysis Report"
        section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        section.footer.paragraphs[0].text = "Calculated from Eye, Blink, Head Pitch and Yawning Data"
        section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def status(self, score, alertness=False):
        if alertness:
            return "Alert" if score >= 85 else "Moderately Alert" if score >= 70 else "Drowsy Risk"
        return "Attentive" if score >= 85 else "Moderately Attentive" if score >= 70 else "Inattentive"

    def prepare_status_data(self):
        self.driver_report["Attentiveness_Score"] = (
            self.driver_report["Eye_Score"] * (0.35 / 0.60)
            + self.driver_report["Head_Score"] * (0.25 / 0.60)).round(2)
        
        self.driver_report["Alertness_Score"] = (
            self.driver_report["Blink_Score"] * (0.25 / 0.40)
            + self.driver_report["Yawning_Score"] * (0.15 / 0.40)).round(2)
        
        self.driver_report["Attentiveness_Status"] = self.driver_report["Attentiveness_Score"].apply(self.status)
        
        self.driver_report["Alertness_Status"] = self.driver_report["Alertness_Score"].apply(
            lambda score: self.status(score, alertness=True))

    def chart_path(self, name):
        return os.path.join(self.chart_folder, name)

    def add_picture(self, path, title, width=5.5):
        if not os.path.exists(path):
            return
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(path, width=Inches(width))
        caption = self.document.add_paragraph(title)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption.runs:
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(9)

    def cover_page(self):
        self.document.add_paragraph("\n\n")
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DRIVER MONITORING SYSTEM")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(31, 78, 120)
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Calculated Driver Attentiveness Report")
        run.bold = True
        run.font.size = Pt(17)
        date = self.document.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.add_run("Generated Using Python\n").bold = True
        date.add_run(datetime.now().strftime("%d %B %Y, %I:%M %p"))
        self.document.add_page_break()

    def executive_summary(self):
        self.document.add_heading("1. Executive Summary", level=1)
        self.document.add_paragraph(
            "The score uses Eye Closure (35%), Blink Rate (25%), Head Pitch (25%) "
            "and Yawning (15%). Results update when the Excel data changes.")
        
    def dashboard(self):
        self.document.add_heading("2. Dashboard Summary", level=1)
        table = self.document.add_table(rows=3, cols=2)
        table.style = "Colorful Grid Accent 1"
        values = [
            ("Total Records", len(self.data)),
            ("Total Drivers", self.data["Driver_ID"].nunique()),
            ("Average Eye Closure", f"{self.data['Eye_Closure_%'].mean():.2f}%"),
            ("Average Blink Rate", f"{self.data['Blink_Rate'].mean():.2f}"),
            ("Average Calculated Score", f"{self.data['Attention_Score'].mean():.2f}"),
            ("Inattentive Records", int((self.data["Attention_Status"] == "Inattentive").sum())),]
        cells = [cell for row in table.rows for cell in row.cells]
        for cell, (label, value) in zip(cells, values):
            cell.text = f"{label}\n{value}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def all_driver_shift_status(self):
        self.document.add_page_break()
        self.document.add_heading("3. All Drivers and All Shifts Status", level=1)
        table = self.document.add_table(rows=1, cols=8)
        table.style = "Colorful Grid Accent 1"
        headings = [
            "Driver", "Shift", "Attentiveness", "Attention Status",
            "Alertness", "Alertness Status", "Overall", "Overall Status"]
        for cell, heading in zip(table.rows[0].cells, headings):
            cell.text = heading
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        order = pd.CategoricalDtype(["Morning", "Afternoon", "Night"], ordered=True)
        report_data = self.driver_report.copy()
        report_data["Shift"] = report_data["Shift"].astype(order)
        report_data = report_data.sort_values(["Driver_Name", "Shift"])

        for driver in report_data["Driver_Name"].unique():
            driver_data = report_data[report_data["Driver_Name"] == driver]
            first_cell = None
            last_cell = None
            for row_number, (_, row) in enumerate(driver_data.iterrows()):
                cells = table.add_row().cells
                cells[0].text = str(driver) if row_number == 0 else ""
                first_cell = cells[0] if row_number == 0 else first_cell
                last_cell = cells[0]
                values = [
                    row["Shift"], f"{row['Attentiveness_Score']:.1f}", row["Attentiveness_Status"],
                    f"{row['Alertness_Score']:.1f}", row["Alertness_Status"],
                    f"{row['Attention_Score']:.1f}", row["Status"],]
                for index, value in enumerate(values, start=1):
                    cells[index].text = str(value)
                for cell in cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged = first_cell.merge(last_cell)
            merged.text = str(driver)
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if merged.paragraphs[0].runs:
                merged.paragraphs[0].runs[0].bold = True

    def driver_summary(self):
        self.document.add_heading("4. Driver Performance Summary", level=1)
        table = self.document.add_table(rows=1, cols=8)
        table.style = "Colorful Grid Accent 1"
        headings = ["Rank", "Driver", "Overall", "Status", "Eye", "Blink", "Head", "Yawn"]
        for cell, heading in zip(table.rows[0].cells, headings):
            cell.text = heading
        for _, row in self.ranking.iterrows():
            values = [
                row["Rank"], row["Driver_Name"], f"{row['Overall_Score']:.2f}", row["Overall_Status"],
                f"{row['Average_Eye_Score']:.1f}", f"{row['Average_Blink_Score']:.1f}",
                f"{row['Average_Head_Score']:.1f}", f"{row['Average_Yawning_Score']:.1f}",]
            for cell, value in zip(table.add_row().cells, values):
                cell.text = str(value)

    def save_chart(self, file_name, dpi=170):
        plt.tight_layout()
        plt.savefig(self.chart_path(file_name), dpi=dpi, bbox_inches="tight")
        plt.close()

    def create_overall_charts(self):
        scores = self.ranking.sort_values("Overall_Score")
        colors = ["seagreen" if s >= 85 else "orange" if s >= 70 else "firebrick" for s in scores["Overall_Score"]]
        plt.figure(figsize=(9, 5))
        bars = plt.barh(scores["Driver_Name"], scores["Overall_Score"], color=colors)
        for bar, score in zip(bars, scores["Overall_Score"]):
            plt.text(score + 1, bar.get_y() + bar.get_height() / 2, f"{score:.1f}", va="center")
        plt.title("Overall Calculated Driver Attentiveness Ranking")
        plt.xlabel("Overall Score")
        plt.xlim(0, 110)
        plt.axvline(85, color="green", linestyle="--", label="Attentive")
        plt.axvline(70, color="orange", linestyle="--", label="Moderate")
        plt.legend()
        self.save_chart("overall_ranking.png")

        comparison = self.ranking.sort_values("Overall_Score", ascending=False)
        colors = ["seagreen" if s >= 85 else "orange" if s >= 70 else "firebrick" for s in comparison["Overall_Score"]]
        plt.figure(figsize=(9, 5.5))
        bars = plt.bar(
            comparison["Driver_Name"], comparison["Overall_Score"],
            color=colors, width=0.65, edgecolor="black", linewidth=0.7,)
        for bar, score in zip(bars, comparison["Overall_Score"]):
            plt.text(
                bar.get_x() + bar.get_width() / 2, bar.get_height() + 1,
                f"{score:.1f}", ha="center", fontweight="bold",)
            
        plt.title("Overall Score Comparison of All Five Drivers")
        plt.xlabel("Driver")
        plt.ylabel("Overall Calculated Score")
        plt.ylim(0, 110)
        plt.axhline(85, color="green", linestyle="--", label="Attentive")
        plt.axhline(70, color="orange", linestyle="--", label="Moderate")
        plt.grid(axis="y", alpha=0.25)
        plt.legend(loc="lower left")
        self.save_chart("all_drivers_overall_bar.png", dpi=180)

        shift_table = self.driver_report.pivot_table(
            index="Driver_Name", columns="Shift", values="Attention_Score",
            aggfunc="mean", observed=False,).reindex(columns=["Morning", "Afternoon", "Night"])
        colors = ["#1f77b4", "#ff7f0e", "#2ca02c", "#9467bd", "#d62728"]
        markers = ["o", "s", "^", "D", "P"]
        plt.figure(figsize=(10, 5.5))
        for index, driver in enumerate(shift_table.index):
            plt.plot(
                ["Morning", "Afternoon", "Night"], shift_table.loc[driver],
                marker=markers[index], linewidth=2.5, color=colors[index], label=driver,)
        plt.title("All Drivers and All Shifts - Attentiveness Comparison")
        plt.xlabel("Shift")
        plt.ylabel("Calculated Attention Score")
        plt.ylim(0, 110)
        plt.axhline(85, color="green", linestyle="--")
        plt.axhline(70, color="orange", linestyle="--")
        plt.grid(axis="y", alpha=0.3)
        plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
        self.save_chart("all_drivers_all_shifts.png", dpi=180)

        components = self.ranking.set_index("Driver_Name")[[
            "Average_Eye_Score", "Average_Blink_Score",
            "Average_Head_Score", "Average_Yawning_Score",]].rename(columns={
            "Average_Eye_Score": "Eye Closure", "Average_Blink_Score": "Blink Rate",
            "Average_Head_Score": "Head Pitch", "Average_Yawning_Score": "Yawning",})
        components.plot(kind="bar", figsize=(11, 5.5),color=["crimson", "seagreen", "mediumpurple", "darkorange"],)
        plt.title("Calculated Component Score Comparison")
        plt.xlabel("Driver")
        plt.ylabel("Component Score")
        plt.ylim(0, 110)
        plt.xticks(rotation=0)
        self.save_chart("component_comparison.png")

    def overall_driver_comparison(self):
        self.document.add_page_break()
        self.document.add_heading("5. Overall Driver and Shift Comparison", level=1)
        best = self.ranking.iloc[0]
        worst = self.ranking.iloc[-1]
        shift_scores = self.driver_report.groupby("Shift", observed=False)["Attention_Score"].mean()
        table = self.document.add_table(rows=4, cols=2)
        table.style = "Colorful Grid Accent 1"
        values = [("Best Overall Driver", f"{best['Driver_Name']} ({best['Overall_Score']:.2f}/100)"),
            ("Best Driver Status", best["Overall_Status"]),
            ("Best Overall Shift", f"{shift_scores.idxmax()} ({shift_scores.max():.2f}/100)"),
            ("Needs Most Attention", f"{worst['Driver_Name']} ({worst['Overall_Score']:.2f}/100)"),]
        for index, (label, value) in enumerate(values):
            table.rows[index].cells[0].text = str(label)
            table.rows[index].cells[1].text = str(value)
        charts = [
            ("all_drivers_overall_bar.png", "Overall Score Bar Comparison of All Five Drivers"),
            ("overall_ranking.png", "Overall Driver Attentiveness Ranking"),
            ("all_drivers_all_shifts.png", "All Drivers and All Shifts in One Chart"),
            ("component_comparison.png", "Eye, Blink, Head Pitch and Yawning Component Scores"),]
        for name, title in charts:
            self.add_picture(self.chart_path(name), title, 5.7)

    def create_driver_charts(self, driver, rows):
        order = pd.CategoricalDtype(["Morning", "Afternoon", "Night"], ordered=True)
        data = rows.copy()
        data["Shift"] = data["Shift"].astype(order)
        data = data.sort_values("Shift")
        name = driver.replace(" ", "_")
        score_path = self.chart_path(f"{name}_score.png")
        component_path = self.chart_path(f"{name}_components.png")
        parameter_path = self.chart_path(f"{name}_parameters.png")

        plt.figure(figsize=(7, 4.5))
        plt.plot(data["Shift"], data["Attention_Score"], marker="o", linewidth=2.5, color="navy")
        plt.title(f"{driver} - Calculated Attention Score")
        plt.ylabel("Score")
        plt.ylim(0, 110)
        plt.grid(True, alpha=0.3)
        self.save_chart(f"{name}_score.png", dpi=160)

        plt.figure(figsize=(7, 4.5))
        for column, label, marker, color in [
            ("Eye_Score", "Eye Score", "o", "crimson"),
            ("Blink_Score", "Blink Score", "s", "seagreen"),
            ("Head_Score", "Head Score", "^", "purple"),
            ("Yawning_Score", "Yawning Score", "D", "darkorange"),]:
            plt.plot(data["Shift"], data[column], marker=marker, linewidth=2, label=label, color=color)
        plt.title(f"{driver} - Calculated Component Scores")
        plt.ylabel("Component Score")
        plt.ylim(0, 110)
        plt.grid(True, alpha=0.3)
        plt.legend()
        self.save_chart(f"{name}_components.png", dpi=160)

        plt.figure(figsize=(7, 4.5))
        for column, label, marker, color in [
            ("Eye_Closure", "Eye Closure %", "o", "red"),
            ("Blink_Rate", "Blink Rate", "s", "green"),
            ("Head_Pitch", "Head Pitch", "^", "purple"),
            ("Yawning", "Yawning Count", "D", "orange"),]:
            plt.plot(data["Shift"], data[column], marker=marker, linewidth=2, label=label, color=color)
        plt.title(f"{driver} - Sensor Parameter Trend")
        plt.ylabel("Measured Value")
        plt.grid(True, alpha=0.3)
        plt.legend()
        self.save_chart(f"{name}_parameters.png", dpi=160)
        return score_path, component_path, parameter_path

    def driver_analysis(self):
        self.document.add_page_break()
        self.document.add_heading("6. Individual Driver Analysis", level=1)
        drivers = self.driver_report["Driver_Name"].unique()
        for number, driver in enumerate(drivers, start=1):
            rows = self.driver_report[self.driver_report["Driver_Name"] == driver].copy()
            self.document.add_heading(f"6.{number} {driver}", level=2)
            table = self.document.add_table(rows=1, cols=9)
            table.style = "Colorful Grid Accent 1"
            headings = ["Shift", "Overall", "Status", "Eye", "Blink", "Head", "Yawn", "Distance", "Duration"]
            for cell, heading in zip(table.rows[0].cells, headings):
                cell.text = heading
            for _, row in rows.iterrows():
                values = [
                    row["Shift"], f"{row['Attention_Score']:.1f}", row["Status"],
                    f"{row['Eye_Score']:.1f}", f"{row['Blink_Score']:.1f}",
                    f"{row['Head_Score']:.1f}", f"{row['Yawning_Score']:.1f}",
                    f"{row['Distance']:.1f}", f"{row['Duration']} min",]
                for cell, value in zip(table.add_row().cells, values):
                    cell.text = str(value)
            paths = self.create_driver_charts(driver, rows)
            titles = [
                "Calculated Attention Score by Shift",
                "Calculated Eye, Blink, Head Pitch and Yawning Component Scores",
                "Measured Eye Closure, Blink Rate, Head Pitch and Yawning",]
            for path, title in zip(paths, titles):
                self.add_picture(path, title, 4.9)
            if number < len(drivers):
                self.document.add_page_break()

    def conclusion(self):
        self.document.add_page_break()
        self.document.add_heading("7. Conclusion",level=1)
        best_driver = self.ranking.iloc[0]
        lowest_driver = self.ranking.iloc[-1]
        shift_scores = self.driver_report.groupby("Shift",observed=False)["Attention_Score"].mean()
        best_shift = shift_scores.idxmax()
        lowest_shift = shift_scores.idxmin()
        average_score = self.ranking["Overall_Score"].mean()
        attentive_count = (self.ranking["Overall_Status"]== "Attentive").sum()
        moderate_count = (self.ranking["Overall_Status"]== "Moderately Attentive").sum()
        inattentive_count = (self.ranking["Overall_Status"]== "Inattentive").sum()
        paragraph = self.document.add_paragraph()
        paragraph.add_run(
            "The calculated analysis produced an average "
         f"driver score of {average_score:.2f}/100.")

        paragraph.add_run(f"{best_driver['Driver_Name']} is the "f"best-performing driver with an overall score "
            f"of {best_driver['Overall_Score']:.2f}/100 "f"and a status of "f"{best_driver['Overall_Status']}.")
        
        paragraph.add_run(f"{lowest_driver['Driver_Name']} recorded the "f"lowest overall score of "
                          f"{lowest_driver['Overall_Score']:.2f}/100 "f"and requires the greatest attention. ")

        paragraph.add_run(f"{best_shift} is the best-performing shift "f"with an average score of "
                        f"{shift_scores.max():.2f}/100, while "f"{lowest_shift} is the lowest-performing shift "
                        f"with an average score of "
                        f"{shift_scores.min():.2f}/100.")
        
        paragraph.add_run(f"Among the {len(self.ranking)} drivers,"f"{attentive_count} are Attentive,"
                          f"{moderate_count} are Moderately Attentive, "
            f"and {inattentive_count} are Inattentive. ")
        
        paragraph.add_run(
        "The results are generated dynamically from "
        "eye closure, blink rate, head pitch and "
        "yawning measurements. Therefore, changing "
        "these values in the Excel dataset and running "
        "the project again will automatically update "
        "the scores, statuses, rankings, charts and "
        "this conclusion.")

    def generate_report(self):
        report_file = "reports/Analysis_Report.docx"
        if os.path.exists(report_file):
            backup = datetime.now().strftime("backup/Report_%Y%m%d_%H%M%S.docx")
            try:
                shutil.copy(report_file, backup)
            except PermissionError:
                pass
        self.create_overall_charts()
        self.cover_page()
        self.executive_summary()
        self.dashboard()
        self.all_driver_shift_status()
        self.driver_summary()
        self.overall_driver_comparison()
        self.driver_analysis()
        self.conclusion()
        try:
            self.document.save(report_file)
        except PermissionError:
            report_file = datetime.now().strftime("reports/Report_%Y%m%d_%H%M%S.docx")
            self.document.save(report_file)
        print("Word Report Generated:", report_file)
        if os.path.exists(self.chart_folder):
            shutil.rmtree(self.chart_folder)
        return report_file