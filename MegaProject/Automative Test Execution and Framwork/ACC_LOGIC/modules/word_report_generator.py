import os
import shutil
from docx import Document

from modules.report_utils import clean_text
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x0F, 0x17, 0x2A)
GREEN = RGBColor(0x15, 0x80, 0x3D)
RED = RGBColor(0xB9, 0x1C, 0x1C)
AMBER = RGBColor(0xB4, 0x53, 0x09)
VIOLET = RGBColor(0x6D, 0x28, 0xD9)
MUTED = RGBColor(0x64, 0x74, 0x8B)

STATUS_COLORS = {
    "Pass": GREEN,
    "Fail": RED,
    "Missing Value Error": AMBER,
    "Invalid Value Error": VIOLET,
}


class WordReportGenerator:

    @staticmethod
    def create_report(
        df,
        pass_count,
        fail_count,
        missing_count,
        invalid_count,
        chart_path,
        output_docx,
        report_title="Test Execution Report",
        report_subtitle="",
    ):
        total_tc = len(df)

        doc = Document()
        WordReportGenerator._set_default_font(doc)

        # ---------------- Title ----------------
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(report_title)
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = NAVY

        if report_subtitle:
            sub_p = doc.add_paragraph()
            sub_run = sub_p.add_run(report_subtitle)
            sub_run.font.size = Pt(12)
            sub_run.font.color.rgb = MUTED

        meta_p = doc.add_paragraph()
        from datetime import datetime
        meta_run = meta_p.add_run(
            f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \u00b7  {total_tc} test case(s) executed"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = MUTED

        WordReportGenerator._add_horizontal_rule(doc)

        # ---------------- Summary stats ----------------
        doc.add_heading("Execution Summary", level=2)
        summary_table = doc.add_table(rows=1, cols=5)
        summary_table.style = "Table Grid"
        summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        headers = ["Total", "Passed", "Failed", "Missing Value Error", "Invalid Value Error"]
        values = [total_tc, pass_count, fail_count, missing_count, invalid_count]
        value_colors = [NAVY, GREEN, RED, AMBER, VIOLET]

        header_cells = summary_table.rows[0].cells
        for i, h in enumerate(headers):
            header_cells[i].text = ""
            run = header_cells[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            WordReportGenerator._shade_cell(header_cells[i], "0F172A")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        value_row = summary_table.add_row().cells
        for i, (v, color) in enumerate(zip(values, value_colors)):
            value_row[i].text = ""
            run = value_row[i].paragraphs[0].add_run(str(v))
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = color

        # ---------------- Pie chart ----------------
        doc.add_heading("Pass / Fail / Missing Distribution", level=2)
        if chart_path and os.path.exists(chart_path):
            chart_p = doc.add_paragraph()
            chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chart_p.add_run().add_picture(chart_path, width=Inches(4.2))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption.add_run("Positive (green) \u00b7 Negative (red) \u00b7 Missing/Invalid Data (yellow)")
            cap_run.font.size = Pt(8.5)
            cap_run.font.color.rgb = MUTED

        # ---------------- Detailed results table ----------------
        doc.add_page_break()
        doc.add_heading("Detailed Test Results", level=2)
        columns = ["TC ID", "Row", "Description", "Action", "Expected", "Actual", "Status", "Remarks", "Missing Fields", "Invalid Fields"]
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        table.autofit = True

        WordReportGenerator._allow_row_break(table.rows[0])
        header_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            header_cells[i].text = ""
            run = header_cells[i].paragraphs[0].add_run(col)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            WordReportGenerator._shade_cell(header_cells[i], "0F172A")

        for _, row in df.iterrows():
            row_obj = table.add_row()
            WordReportGenerator._allow_row_break(row_obj)
            cells = row_obj.cells
            values = [
                str(row["TC_ID"]),
                str(row["Excel_Row"]),
                str(row["Test_Description"]) or "(blank)",
                str(row["User_Action"]) or "(blank)",
                str(row["Expected_Result"]) if str(row["Expected_Result"]).strip() else "(blank)",
                str(row["Actual_Result"]),
                str(row["Status"]),
                str(row["Remarks"]),
                row["Missing_Fields"] if row["Missing_Fields"] else "\u2014",
                row["Invalid_Fields"] if row["Invalid_Fields"] else "\u2014",
            ]
            for i, val in enumerate(values):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(val)
                run.font.size = Pt(7.5)
                if i == 6:  # Status column
                    run.bold = True
                    run.font.color.rgb = STATUS_COLORS.get(row["Status"], NAVY)
                if row["Status"] != "Pass":
                    WordReportGenerator._shade_cell(cells[i], "FBFBFB")

        # ---------------- Failed / Missing / Invalid sections ----------------
        WordReportGenerator._add_issue_section(
            doc, "Failed Test Cases", df[df["Status"] == "Fail"], RED,
            lambda r: (
                f"Expected: {r['Expected_Result']}   Actual: {r['Actual_Result']}\nRemarks: {r['Remarks']}"
                + (f"\nReason: {r['Mismatch_Reason']}" if r.get("Mismatch_Reason") else "")
            ),
            "No failed test cases.",
        )
        WordReportGenerator._add_issue_section(
            doc, "Missing Value Errors", df[df["Status"] == "Missing Value Error"], AMBER,
            lambda r: f"Missing: {r['Missing_Fields'] if r['Missing_Fields'] else 'Expected Result'}\nDetails: {r['Remarks']}",
            "No missing value errors - every required cell was filled in.",
        )
        WordReportGenerator._add_issue_section(
            doc, "Invalid Value Errors", df[df["Status"] == "Invalid Value Error"], VIOLET,
            lambda r: f"Invalid: {r['Invalid_Fields']}\nDetails: {r['Remarks']}",
            "No invalid value errors - every recognized cell matched its expected format.",
        )

        output_dir = os.path.dirname(os.path.abspath(output_docx)) or "."
        os.makedirs(output_dir, exist_ok=True)
        temp_dir = os.path.join(os.path.dirname(os.path.abspath(output_docx)), "_tmp")
        os.makedirs(temp_dir, exist_ok=True)
        temp_docx = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(output_docx))[0]}_{os.getpid()}.docx")
        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except OSError:
                pass
        doc.save(temp_docx)
        final_docx = os.path.abspath(output_docx)
        if os.path.exists(final_docx):
            try:
                os.chmod(final_docx, 0o666)
            except OSError:
                pass
        shutil.copy2(temp_docx, final_docx)
        try:
            os.remove(temp_docx)
        except OSError:
            pass
        print(f"[WordReportGenerator] Word report saved to '{output_docx}'")

    # -----------------------------------------------------------
    # Helpers
    # -----------------------------------------------------------

    @staticmethod
    def _set_default_font(doc):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    @staticmethod
    def _add_horizontal_rule(doc):
        p = doc.add_paragraph()
        p_fmt = p.paragraph_format
        p_fmt.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D7DEE6")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _shade_cell(cell, hex_color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    @staticmethod
    def _allow_row_break(row):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:canBreak")) is None:
            tr_pr.append(OxmlElement("w:canBreak"))

    @staticmethod
    def _add_issue_section(doc, heading, subset_df, accent_color, body_fn, empty_message):
        doc.add_heading(heading, level=2)
        if len(subset_df) == 0:
            p = doc.add_paragraph(empty_message)
            p.runs[0].font.color.rgb = MUTED
            return
        for _, row in subset_df.iterrows():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            id_run = p.add_run(f"{row['TC_ID']} \u00b7 Row {row['Excel_Row']}")
            id_run.bold = True
            id_run.font.color.rgb = accent_color
            id_run.font.size = Pt(9.5)

            body_p = doc.add_paragraph()
            body_p.paragraph_format.space_after = Pt(10)
            body_run = body_p.add_run(body_fn(row))
            body_run.font.size = Pt(9)
