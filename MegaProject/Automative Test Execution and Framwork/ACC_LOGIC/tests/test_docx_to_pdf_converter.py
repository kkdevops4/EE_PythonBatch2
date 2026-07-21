import os
import tempfile
import unittest
from unittest import mock

from docx import Document
from pypdf import PdfReader

from modules.docx_to_pdf_converter import DocxToPDFConverter


class DocxToPDFConverterTests(unittest.TestCase):
    def test_uses_docx2pdf_fallback_when_libreoffice_is_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            source_docx = os.path.join(tmpdir, "ACC_Report.docx")
            target_pdf = os.path.join(tmpdir, "ACC_Report.pdf")
            with open(source_docx, "wb") as handle:
                handle.write(b"fake docx")

            with mock.patch("modules.docx_to_pdf_converter.shutil.which", side_effect=lambda name: None):
                def fake_convert(src, dst):
                    with open(dst, "wb") as handle:
                        handle.write(b"pdf")

                with mock.patch("modules.docx_to_pdf_converter.docx2pdf.convert", side_effect=fake_convert):
                    result = DocxToPDFConverter.convert(source_docx, target_pdf)

            self.assertTrue(result)
            self.assertTrue(os.path.exists(target_pdf))

    def test_creates_pdf_from_docx_when_no_converter_is_available(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            source_docx = os.path.join(tmpdir, "ACC_Report.docx")
            target_pdf = os.path.join(tmpdir, "ACC_Report.pdf")

            document = Document()
            document.add_paragraph("ACC Test Report")
            document.add_paragraph("Summary: 10 passed, 1 failed")
            for index in range(120):
                document.add_paragraph(f"Detail paragraph {index + 1}: This is a long report section that should span multiple PDF pages when converted.")
            document.save(source_docx)

            with mock.patch("modules.docx_to_pdf_converter.shutil.which", side_effect=lambda name: None):
                with mock.patch("modules.docx_to_pdf_converter.docx2pdf.convert", side_effect=RuntimeError("not available")):
                    result = DocxToPDFConverter.convert(source_docx, target_pdf)

            self.assertTrue(result)
            self.assertTrue(os.path.exists(target_pdf))

            reader = PdfReader(target_pdf)
            self.assertGreater(len(reader.pages), 1)


if __name__ == "__main__":
    unittest.main()
